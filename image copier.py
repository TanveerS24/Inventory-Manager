from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT as WDO
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WDPA
import os

# === CONFIGURATION ===
template_path = r"C:\Users\Tanveer\Python\Inventory Manager" #insert template document path here
source_folder = r"C:\Users\Tanveer\Pictures\Screenshots" #insert source image folder path here
output_file = "Photo_gallery.docx" #output file
images_per_page = 8
images_per_row = 4
image_width = Cm(5.85)
image_height = Cm(6.11)

# === STEP 1: Load the template first page ===
doc = Document(template_path)

# === STEP 2: Add a new section (landscape) for photo pages ===
landscape_section = doc.add_section(start_type=1)  # New page section
landscape_section.orientation = WDO.LANDSCAPE
landscape_section.page_width, landscape_section.page_height = (
    landscape_section.page_height, landscape_section.page_width
)

# Just to be safe, enforce landscape on all sections after the first
for idx, sec in enumerate(doc.sections):
    if idx != 0:
        sec.orientation = WDO.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width

# === STEP 3: Get all image files sorted ===
image_files = sorted([
    f for f in os.listdir(source_folder)
    if f.lower().endswith(('.png', '.jpg', '.jpeg'))
])

photo_counter = 1

# === STEP 4: Insert images in sets of 8 per page ===
for start in range(0, len(image_files), images_per_page):
    if start != 0:
        doc.add_page_break()

    # PHOTO INDEX heading
    heading = doc.add_paragraph()
    heading.alignment = WDPA.LEFT
    heading.paragraph_format.space_after = Pt(12)  # space below title
    heading_run = heading.add_run("PHOTO INDEX")
    heading_run.bold = True
    heading_run.font.size = Pt(20)

    # 4-row table: 2 image rows and 2 caption rows
    table = doc.add_table(rows=4, cols=images_per_row)
    table.autofit = True

    for i in range(images_per_page):
        idx = start + i
        if idx >= len(image_files):
            break

        row = (i // images_per_row) * 2  # row 0 or 2
        col = i % images_per_row
        img_path = os.path.join(source_folder, image_files[idx])

        # Insert image
        img_cell = table.cell(row, col)
        img_para = img_cell.paragraphs[0]
        img_para.paragraph_format.space_after = Pt(0)
        img_para.alignment = WDPA.CENTER
        run = img_para.add_run()
        run.add_picture(img_path, width=image_width, height=image_height)

        # Insert caption
        caption_cell = table.cell(row + 1, col)
        caption_para = caption_cell.paragraphs[0]
        caption_para.alignment = WDPA.CENTER
        caption_para.paragraph_format.space_before = Pt(10)
        caption_para.paragraph_format.space_after = Pt(10)
        caption_run = caption_para.add_run(f"Photo {photo_counter:03d}")
        caption_run.font.size = Pt(12)

        photo_counter += 1

# === STEP 5: Save final document ===
if os.path.exists(output_file):
    os.remove(output_file)

doc.save(output_file)
print(f"✅ Done! Document saved as '{output_file}' with {photo_counter - 1} photos.")
