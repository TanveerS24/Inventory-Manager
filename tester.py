import sqlite3
from datetime import datetime
import time
import tkinter as tk
from tkinter import messagebox, ttk, filedialog
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Cm, Pt, Inches, Mm, RGBColor
from docx.enum.section import WD_ORIENT as WDO
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WDPA
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx2pdf import convert
from PyPDF2 import PdfMerger
from pdf2docx import Converter
import os, sys, platform, subprocess, psutil, tempfile, shutil, traceback

# ---------------- CONFIG ----------------
CLERK_OPTIONS = ["Tom Tyrrel", "Kevin Crack", "Bill West"]
STATUS_OPTIONS = ["Inspected", "Audio Recorded"]
APP_BG = "#9ECAD6"
ACTION_BG = "#DFF6FF"
BTN_BG = "#748DAE"
BTN_FG = "#2C3E50"

if getattr(sys, 'frozen', False):
    BASE_PATH = os.path.dirname(sys.executable)
else:
    BASE_PATH = os.path.dirname(os.path.abspath(__file__))

DB_PATH = os.path.join(BASE_PATH, "inventory.db")

# ---------------- Utilities ----------------
def connect_db():
    return sqlite3.connect(DB_PATH)

def create_table():
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS property_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        date TEXT,
        clerk TEXT,
        property_address TEXT,
        client TEXT,
        inv_type TEXT,
        status TEXT
    )
    """)
    conn.commit()
    conn.close()

def set_doc_landscape(doc: Document):
    for section in doc.sections:
        section.orientation = WDO.LANDSCAPE
        section.page_width, section.page_height = Mm(297), Mm(210)
        # reasonable margins
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

def add_table_borders(table):
    tbl_el = table._tbl
    tbl_pr = tbl_el.tblPr or OxmlElement('w:tblPr')
    tbl_el.insert(0, tbl_pr)
    tbl_borders = parse_xml(
        r'<w:tblBorders %s>'
        r'<w:top w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:left w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:right w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    for old in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(old)
    tbl_pr.append(tbl_borders)

def safe_convert_docx_to_pdf(input_path, output_path):
    """
    Use docx2pdf.convert to convert a single file.
    docx2pdf will call MS Word COM on Windows. We run sequentially and avoid forcing word termination.
    """
    try:
        convert(input_path, output_path)  # docx2pdf handles same-path mapping
    except Exception as e:
        raise RuntimeError(f"docx->pdf conversion failed for {input_path}: {e}")

def ensure_word_closed_gracefully(wait_seconds=2):
    """
    Optionally try to close any leftover WINWORD.EXE instances politely.
    We avoid killing forcibly. This is only used sparingly.
    """
    if platform.system() != "Windows":
        return
    for proc in psutil.process_iter(['pid', 'name']):
        try:
            name = proc.info.get('name', '') or ''
            if 'WINWORD.EXE' in name.upper():
                try:
                    proc.terminate()
                except Exception:
                    pass
        except Exception:
            pass
    # small wait
    time.sleep(wait_seconds)

# ---------------- Document Builders ----------------
def generate_template_docx(address, client, date_str, logo_path=None, out_path=None):
    """Create and return path to template DOCX (landscape)."""
    if out_path is None:
        out_path = os.path.join(tempfile.gettempdir(), f"template_{int(time.time())}.docx")
    doc = Document()
    set_doc_landscape(doc)

    # logo
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WDPA.LEFT
        r = p.add_run()
        r.add_picture(logo_path, width=Cm(12.38), height=Cm(2.9))

    # placeholders
    for text in [
        f"Property Address: {address}",
        f"On behalf of:     {client}",
        f"Date:             {date_str}"
    ]:
        para = doc.add_paragraph(text)
        para.runs[0].bold = True

    doc.add_paragraph()
    ph = doc.add_paragraph("Additional Notes")
    ph.runs[0].bold = True

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(25.46)
    tbl.rows[0].height = Cm(5.19)
    add_table_borders(tbl)
    tbl.rows[0].cells[0].paragraphs[0].text = "Property Photos Link"

    doc.add_paragraph()
    f1 = doc.add_paragraph()
    run1 = f1.add_run("Inventory House ")
    run1.font.color.rgb = RGBColor(255, 0, 0)
    run1.font.size = Pt(12)
    run1.bold = True
    run2 = f1.add_run("T: 08700 336969 ")
    run2.font.size = Pt(12)

    email = "info@inventoryhouse.co.uk"
    website = "www.inventoryhouse.co.uk"
    # add email / website plainly (hyperlink creation is more verbose; plain text is fine)
    f1.add_run(f"  {email}  ")
    f1.add_run(f"  {website}")

    f2 = doc.add_paragraph(
        "Head Office: 3 County Gate London SE9 3UB.\n"
        "Inventory House Limited. Registered in England & Wales Company No. 5250554"
    )
    f2.runs[0].font.size = Pt(12)
    f2.runs[0].bold = True

    doc.save(out_path)
    return out_path

def force_docx_to_landscape_and_save(input_path, out_path=None):
    """
    Load a docx, force landscape orientation, save to out_path.
    Returns out_path.
    """
    if out_path is None:
        out_path = os.path.join(tempfile.gettempdir(), f"land_{int(time.time())}.docx")
    doc = Document(input_path)
    set_doc_landscape(doc)
    doc.save(out_path)
    return out_path

def build_photo_index_docx(source_folder, out_path=None, images_per_page=8, images_per_row=4):
    """
    Creates a DOCX with images arranged in a grid and captions.
    Ensures landscape sections.
    Returns path to the generated DOCX.
    """
    if out_path is None:
        out_path = os.path.join(tempfile.gettempdir(), f"photos_{int(time.time())}.docx")
    image_files = sorted([f for f in os.listdir(source_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg'))])
    if not image_files:
        raise ValueError("No image files found in selected folder.")

    doc = Document()
    set_doc_landscape(doc)

    image_width = Cm(5.85)
    image_height = Cm(6)
    photo_counter = 1

    for start in range(0, len(image_files), images_per_page):
        heading = doc.add_paragraph()
        heading.alignment = WDPA.LEFT
        heading.paragraph_format.space_after = Pt(8)
        heading_run = heading.add_run("PHOTO INDEX")
        heading_run.bold = True
        heading_run.font.size = Pt(20)

        # create table with rows = images_per_row*2 (image row + caption row pairs) but we will create 'rows = images_per_page//images_per_row * 2'
        rows_needed = (min(images_per_page, len(image_files) - start) + images_per_row - 1) // images_per_row
        table = doc.add_table(rows=rows_needed * 2, cols=images_per_row)
        table.autofit = True
        add_table_borders(table)

        for i in range(min(images_per_page, len(image_files) - start)):
            global_index = start + i
            row = (i // images_per_row) * 2
            col = i % images_per_row
            img_path = os.path.join(source_folder, image_files[global_index])

            img_cell = table.cell(row, col)
            img_para = img_cell.paragraphs[0]
            img_para.paragraph_format.space_after = Pt(0)
            img_para.alignment = WDPA.CENTER
            run = img_para.add_run()
            run.add_picture(img_path, width=image_width, height=image_height)

            cap_cell = table.cell(row + 1, col)
            cap_para = cap_cell.paragraphs[0]
            cap_para.alignment = WDPA.CENTER
            cap_para.paragraph_format.space_before = Pt(7)
            cap_run = cap_para.add_run(f"Photo {photo_counter:03d}")
            cap_run.font.size = Pt(12)
            photo_counter += 1

        # page break after each page block (except maybe last)
        doc.add_page_break()

    doc.save(out_path)
    return out_path

# ---------------- GUI / App Functions ----------------
def add_record_popup():
    popup = tk.Toplevel(root)
    popup.title("Add New Record")

    tk.Label(popup, text="Clerk").grid(row=0, column=0)
    clerk_cb = ttk.Combobox(popup, values=CLERK_OPTIONS, state="readonly")
    clerk_cb.grid(row=0, column=1)

    tk.Label(popup, text="Address").grid(row=1, column=0)
    addr_entry = tk.Entry(popup)
    addr_entry.grid(row=1, column=1)

    tk.Label(popup, text="Client").grid(row=2, column=0)
    client_entry = tk.Entry(popup)
    client_entry.grid(row=2, column=1)

    tk.Label(popup, text="Inventory Type").grid(row=3, column=0)
    inv_entry = tk.Entry(popup)
    inv_entry.grid(row=3, column=1)

    tk.Label(popup, text="Status").grid(row=4, column=0)
    status_cb = ttk.Combobox(popup, values=STATUS_OPTIONS, state="readonly")
    status_cb.grid(row=4, column=1)

    def submit():
        values = [clerk_cb.get(), addr_entry.get(), client_entry.get(), inv_entry.get(), status_cb.get()]
        if not all(values):
            messagebox.showwarning("Input Error", "Please fill in all fields.")
            return
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
        INSERT INTO property_records (date, clerk, property_address, client, inv_type, status)
        VALUES (?, ?, ?, ?, ?, ?)
        """, (datetime.now().strftime("%Y-%m-%d"), *values))
        conn.commit()
        conn.close()
        popup.destroy()
        fetch_all()
        messagebox.showinfo("Success", "Record added successfully!")

    tk.Button(popup, text="Add Record", command=submit).grid(row=5, columnspan=2, pady=10)

def search_popup():
    popup = tk.Toplevel(root)
    popup.title("Search Records")

    tk.Label(popup, text="Client").grid(row=0, column=0)
    client_entry = tk.Entry(popup)
    client_entry.grid(row=0, column=1)

    tk.Label(popup, text="Clerk").grid(row=1, column=0)
    clerk_cb = ttk.Combobox(popup, values=CLERK_OPTIONS, state="readonly")
    clerk_cb.grid(row=1, column=1)

    tk.Label(popup, text="Address").grid(row=2, column=0)
    address_entry = tk.Entry(popup)
    address_entry.grid(row=2, column=1)

    def search():
        client = client_entry.get()
        clerk = clerk_cb.get()
        address = address_entry.get()
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
        SELECT * FROM property_records
        WHERE client LIKE ? AND clerk LIKE ? AND property_address LIKE ?
        """, (f"%{client}%", f"%{clerk}%", f"%{address}%"))
        rows = cursor.fetchall()
        conn.close()
        formatted = [(i+1, row[0], datetime.strptime(row[1], "%Y-%m-%d").strftime("%d-%m-%Y"), *row[2:]) for i, row in enumerate(rows)]
        update_tree(formatted)
        btn_clear_filters.config(state="normal")
        popup.destroy()

    tk.Button(popup, text="Search", command=search).grid(row=3, columnspan=2, pady=10)

def fetch_all():
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM property_records")
    rows = cursor.fetchall()
    conn.close()
    formatted = [(i+1, row[0], datetime.strptime(row[1], "%Y-%m-%d").strftime("%d-%m-%Y"), *row[2:]) for i, row in enumerate(rows)]
    update_tree(formatted)

def delete_record(record_id):
    if messagebox.askyesno("Confirm", "Delete this record?"):
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM property_records WHERE id = ?", (record_id,))
        conn.commit()
        conn.close()
        fetch_all()
        messagebox.showinfo("Deleted", "Record deleted successfully.")

def open_edit_popup(record):
    popup = tk.Toplevel(root)
    popup.title("Edit Record")

    is_completed = record[7].strip().lower() == "completed"

    tk.Label(popup, text="Clerk").grid(row=0, column=0, padx=5, pady=5)
    clerk_cb = ttk.Combobox(popup, values=CLERK_OPTIONS, state="readonly")
    clerk_cb.set(record[3])
    clerk_cb.grid(row=0, column=1, padx=5, pady=5)

    tk.Label(popup, text="Address").grid(row=1, column=0, padx=5, pady=5)
    address_entry = tk.Entry(popup)
    address_entry.insert(0, record[4])
    address_entry.grid(row=1, column=1, padx=5, pady=5)

    tk.Label(popup, text="Client").grid(row=2, column=0, padx=5, pady=5)
    client_entry = tk.Entry(popup)
    client_entry.insert(0, record[5])
    client_entry.grid(row=2, column=1, padx=5, pady=5)

    tk.Label(popup, text="Inventory Type").grid(row=3, column=0, padx=5, pady=5)
    inv_entry = tk.Entry(popup)
    inv_entry.insert(0, record[6])
    inv_entry.grid(row=3, column=1, padx=5, pady=5)

    tk.Label(popup, text="Status").grid(row=4, column=0, padx=5, pady=5)
    status_cb = ttk.Combobox(popup, values=STATUS_OPTIONS, state="readonly" if not is_completed else "disabled")
    status_cb.set(record[7])
    status_cb.grid(row=4, column=1, padx=5, pady=5)

    def save_changes():
        clerk = clerk_cb.get()
        address = address_entry.get().strip()
        client = client_entry.get().strip()
        inv_type = inv_entry.get().strip()
        status = record[7] if is_completed else status_cb.get()

        if not all([clerk, address, client, inv_type, status]):
            messagebox.showwarning("Input Error", "Please fill in all fields.")
            return

        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE property_records 
            SET clerk = ?, property_address = ?, client = ?, inv_type = ?, status = ?
            WHERE id = ?
        """, (clerk, address, client, inv_type, status, record[1]))
        conn.commit()
        conn.close()

        popup.destroy()
        fetch_all()
        messagebox.showinfo("Updated", "Record updated successfully.")

    tk.Button(popup, text="Save", command=save_changes).grid(row=5, columnspan=2, pady=10)

# ---------------- Paste Photos Workflow ----------------
def paste_photos(record_id):
    """
    Full workflow:
    1. Ask for middle doc -> force landscape and convert to PDF
    2. Generate template docx -> convert to PDF
    3. Ask for image folder -> create photo index docx -> convert to PDF
    4. Merge PDFs in order (template, middle, images)
    5. Convert final merged PDF to DOCX
    6. Update DB status to Completed and open final files
    """
    try:
        # 1) Middle doc selection & immediate conversion
        middle_doc = filedialog.askopenfilename(title="Select middle Word file (audio transcribed doc)",
                                                filetypes=[("Word Documents", "*.docx;*.doc"), ("All Files", "*.*")])
        if not middle_doc:
            messagebox.showinfo("Cancelled", "No middle document selected.")
            return

        # force landscape and save a temp copy
        middle_land_doc = os.path.join(tempfile.gettempdir(), f"middle_land_{int(time.time())}.docx")
        force_docx_to_landscape_and_save(middle_doc, middle_land_doc)
        middle_pdf = os.path.join(tempfile.gettempdir(), f"middle_{int(time.time())}.pdf")

        # convert middle docx -> pdf (sequential)
        safe_convert_docx_to_pdf(middle_land_doc, middle_pdf)

        # 2) generate template from DB record info and convert
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("SELECT property_address, client, date FROM property_records WHERE id = ?", (record_id,))
        res = cursor.fetchone()
        conn.close()
        if not res:
            messagebox.showerror("Error", "Record not found in database.")
            return
        addr, client, date_val = res
        # date_val should be YYYY-MM-DD in DB; keep displayed date in dd-mm-YYYY for template
        try:
            display_date = datetime.strptime(date_val, "%Y-%m-%d").strftime("%d-%m-%Y")
        except Exception:
            display_date = date_val

        logo = os.path.join(BASE_PATH, "image 1.png")
        template_doc = generate_template_docx(addr, client, display_date, logo_path=logo)
        template_land_doc = template_doc  # already landscape in generator
        template_pdf = os.path.join(tempfile.gettempdir(), f"template_{int(time.time())}.pdf")
        safe_convert_docx_to_pdf(template_land_doc, template_pdf)

        # 3) select image folder -> build photo index docx -> convert to PDF
        source_folder = filedialog.askdirectory(title="Select Image Folder")
        if not source_folder:
            messagebox.showinfo("Cancelled", "No image folder selected.")
            return

        master_doc = build_photo_index_docx(source_folder)
        # master_doc is already landscape
        images_pdf = os.path.join(tempfile.gettempdir(), f"images_{int(time.time())}.pdf")
        safe_convert_docx_to_pdf(master_doc, images_pdf)

        # 4) Merge PDFs in chosen order - we will merge template -> middle -> images
        merged_pdf = os.path.join(BASE_PATH, f"final_{record_id}_{int(time.time())}.pdf")
        merger = PdfMerger()
        for pdf in (template_pdf, middle_pdf, images_pdf):
            merger.append(pdf)
        merger.write(merged_pdf)
        merger.close()

        # 5) Convert merged PDF back to DOCX
        final_docx = os.path.join(BASE_PATH, f"final_{record_id}_{int(time.time())}.docx")
        cv = Converter(merged_pdf)
        cv.convert(final_docx, start=0, end=None)
        cv.close()

        # 6) Update DB status and refresh view
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("UPDATE property_records SET status = ? WHERE id = ?", ("Completed", record_id))
        conn.commit()
        conn.close()
        fetch_all()

        # 7) Show success and open files
        messagebox.showinfo("Success", f"Final PDF saved:\n{merged_pdf}\nFinal DOCX saved:\n{final_docx}")
        if platform.system() == "Windows":
            os.startfile(merged_pdf)
        elif platform.system() == "Darwin":
            subprocess.call(["open", merged_pdf])
        else:
            subprocess.call(["xdg-open", merged_pdf])

    except Exception as e:
        tb = traceback.format_exc()
        messagebox.showerror("Error during paste photos", f"{str(e)}\n\n{tb}")

# ---------------- GUI Setup ----------------
root = tk.Tk()
root.title("InventoryHouse")
root.configure(bg=APP_BG)
root.geometry("1150x700")

# header logo
logo_path = os.path.join(BASE_PATH, "logo.png")
if os.path.exists(logo_path):
    try:
        img = Image.open(logo_path)
        photo = ImageTk.PhotoImage(img)
        title_label = tk.Label(root, bg=APP_BG, image=photo)
        title_label.image = photo
        title_label.pack(pady=10)
    except Exception:
        title_label = tk.Label(root, text="InventoryHouse", bg=APP_BG, font=("Segoe UI", 20, "bold"))
        title_label.pack(pady=10)
else:
    title_label = tk.Label(root, text="InventoryHouse", bg=APP_BG, font=("Segoe UI", 20, "bold"))
    title_label.pack(pady=10)

style = ttk.Style()
style.theme_use('clam')
style.configure("Treeview.Heading", background="#98A1BC", foreground="#555879", font=('Segoe UI', 15, 'bold'))

tree_frame = tk.Frame(root)
tree_frame.pack(fill="both", expand=True, padx=10, pady=5)
tree_y_scroll = tk.Scrollbar(tree_frame, orient="vertical")
tree_y_scroll.pack(side="right", fill="y")
tree_x_scroll = tk.Scrollbar(tree_frame, orient="horizontal")
tree_x_scroll.pack(side="bottom", fill="x")

tree = ttk.Treeview(tree_frame, columns=("No", "ID", "Date", "Clerk", "Address", "Client", "Type", "Status"),
                    show="headings", yscrollcommand=tree_y_scroll.set, xscrollcommand=tree_x_scroll.set)
tree.tag_configure('oddrow', background='white', font=('Segoe UI', 12))
tree.tag_configure('evenrow', background='lightgray', font=('Segoe UI', 12))

def update_tree(rows):
    for item in tree.get_children():
        tree.delete(item)
    for i, row in enumerate(rows):
        tag = 'evenrow' if i % 2 == 0 else 'oddrow'
        tree.insert("", tk.END, values=row, tags=(tag,))
    reset_action_buttons()

tree_y_scroll.config(command=tree.yview)
tree_x_scroll.config(command=tree.xview)
for col in tree["columns"]:
    tree.heading(col, text=col)
    if col == "ID":
        tree.column(col, width=0, stretch=False)
    else:
        tree.column(col, width=120, anchor="center")
tree.pack(fill="both", expand=True)
tree.column("No", width=30, anchor="center")
tree.column("ID", width=0, stretch=False)
tree.column("Date", width=70, anchor="center")
tree.column("Clerk", width=80, anchor="center")
tree.column("Address", width=350, anchor="center", stretch=True)
tree.column("Client", width=100, anchor="center")
tree.column("Type", width=70, anchor="center")
tree.column("Status", width=90, anchor="center")

# Actions
action_frame = tk.LabelFrame(root, text="Table Actions", bg=ACTION_BG, fg="#3C5B6F", font=("Arial", 17, "bold"))
action_frame.pack(fill="x", padx=10, pady=5)
tk.Button(action_frame, text="Add Record", bg=BTN_BG, fg=BTN_FG, font=("Arial", 14, "bold"), command=add_record_popup).pack(side="left", padx=5)
tk.Button(action_frame, text="Search", bg=BTN_BG, fg=BTN_FG, font=("Arial", 14, "bold"), command=search_popup).pack(side="left", padx=5)
btn_clear_filters = tk.Button(action_frame, text="Clear Filters", state="disabled", bg=BTN_BG, fg=BTN_FG, font=("Arial", 14, "bold"), command=lambda: clear_filters())
btn_clear_filters.pack(side="left", padx=5)

row_action_frame = tk.LabelFrame(root, text="Selected Row Actions", bg=ACTION_BG, fg="#3C5B6F", font=("Arial", 17, "bold"))
row_action_frame.pack(fill="x", padx=10, pady=5)
btn_edit = tk.Button(row_action_frame, text="Edit", state="disabled", bg=BTN_BG, fg=BTN_FG, font=("Arial", 14, "bold"))
btn_edit.pack(side="left", padx=5)
btn_delete = tk.Button(row_action_frame, text="Delete", state="disabled", bg=BTN_BG, fg=BTN_FG, font=("Arial", 14, "bold"))
btn_delete.pack(side="left", padx=5)
btn_photos = tk.Button(row_action_frame, text="Paste Photos", state="disabled", bg=BTN_BG, fg=BTN_FG, font=("Arial", 14, "bold"))
btn_photos.pack(side="left", padx=5)

selected_record = None
def on_row_select(event):
    global selected_record
    selected_item = tree.selection()
    if selected_item:
        selected_record = tree.item(selected_item[0])['values']
        btn_edit.config(state="normal", command=lambda: open_edit_popup(selected_record))
        if str(selected_record[7]).strip().lower() == "completed":
            btn_photos.config(state="disabled")
        else:
            btn_photos.config(state="normal", command=lambda: paste_photos(selected_record[1]))
        btn_delete.config(state="normal", command=lambda: delete_record(selected_record[1]))
    else:
        reset_action_buttons()

def reset_action_buttons():
    btn_edit.config(state="disabled")
    btn_delete.config(state="disabled")
    btn_photos.config(state="disabled")

tree.bind("<<TreeviewSelect>>", lambda e: on_row_select(e))

def clear_filters():
    fetch_all()
    btn_clear_filters.config(state="disabled")

# ---------------- Start App ----------------
create_table()
fetch_all()
root.mainloop()
