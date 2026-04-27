"""Document generation service - Reuses original InventoryHouse logic"""
import os
import time
import tempfile
import shutil
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Cm, Pt, Inches, Mm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image as PILImage
from docxcompose.composer import Composer

from app.core.config import settings
from app.models.record import PropertyRecord


class DocumentService:
    """Service for document generation using original logic"""
    
    def __init__(self):
        self.company_info = {
            'name': settings.COMPANY_NAME,
            'phone': settings.COMPANY_PHONE,
            'email': settings.COMPANY_EMAIL,
            'website': settings.COMPANY_WEBSITE,
            'address': settings.COMPANY_ADDRESS,
            'registration': settings.COMPANY_REGISTRATION
        }
    
    def set_doc_landscape(self, doc: Document):
        """Force landscape orientation - from original code"""
        for section in doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = Mm(297), Mm(210)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
    
    def add_table_borders(self, table):
        """Add table borders - from original code"""
        tbl_el = table._tbl
        tbl_pr = tbl_el.tblPr or OxmlElement('w:tblPr')
        tbl_el.insert(0, tbl_pr)
        
        tbl_borders = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
            r'<w:left w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
            r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
            r'<w:right w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
            r'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
            r'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        
        for old in tbl_pr.findall(qn('w:tblBorders')):
            tbl_pr.remove(old)
        tbl_pr.append(tbl_borders)
    
    def generate_template_docx(self, record: PropertyRecord, out_path: str) -> str:
        """Create template DOCX - from original code"""
        doc = Document()
        self.set_doc_landscape(doc)
        
        # Logo
        logo_path = os.path.join(settings.ASSETS_PATH, "image 1.png")
        if os.path.exists(logo_path):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            r = p.add_run()
            r.add_picture(logo_path, width=Cm(12.38), height=Cm(2.9))
        
        # Format date
        try:
            display_date = datetime.strptime(record.date, "%Y-%m-%d").strftime("%d-%m-%Y")
        except:
            display_date = record.date
        
        # Property info
        info_lines = [
            f"Property Address: {record.property_address}",
            f"On behalf of:     {record.client}",
            f"Date:             {display_date}"
        ]
        
        for text in info_lines:
            para = doc.add_paragraph(text)
            para.runs[0].bold = True
        
        doc.add_paragraph()
        
        # Notes table
        ph = doc.add_paragraph("Additional Notes")
        ph.runs[0].bold = True
        
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.columns[0].width = Cm(25.46)
        tbl.rows[0].height = Cm(5.19)
        self.add_table_borders(tbl)
        tbl.rows[0].cells[0].paragraphs[0].text = "Property Photos Link"
        
        doc.add_paragraph()
        
        # Footer
        f1 = doc.add_paragraph()
        run1 = f1.add_run(f"{self.company_info['name']} ")
        run1.font.color.rgb = RGBColor(255, 0, 0)
        run1.font.size = Pt(12)
        run1.bold = True
        
        run2 = f1.add_run(f"T: {self.company_info['phone']} ")
        run2.font.size = Pt(12)
        
        f1.add_run(f"  {self.company_info['email']}  ")
        f1.add_run(f"  {self.company_info['website']}")
        
        f2 = doc.add_paragraph(
            f"{self.company_info['address']}\n"
            f"{self.company_info['registration']}"
        )
        f2.runs[0].font.size = Pt(12)
        f2.runs[0].bold = True
        
        doc.save(out_path)
        return out_path
    
    def force_docx_to_landscape_and_save(self, input_path: str, out_path: str) -> str:
        """Force landscape and save - from original code"""
        doc = Document(input_path)
        self.set_doc_landscape(doc)
        doc.save(out_path)
        return out_path
    
    def build_photo_index_docx(
        self, 
        source_folder: str, 
        out_path: str,
        images_per_page: int = 8,
        images_per_row: int = 4
    ) -> str:
        """Create photo index DOCX - from original code"""
        image_files = sorted([
            f for f in os.listdir(source_folder)
            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
        ])
        
        if not image_files:
            raise ValueError("No image files found in selected folder.")
        
        doc = Document()
        self.set_doc_landscape(doc)
        
        image_width = Cm(settings.PHOTO_WIDTH_CM)
        image_height = Cm(settings.PHOTO_HEIGHT_CM)
        photo_counter = 1
        
        for start in range(0, len(image_files), images_per_page):
            # Heading
            heading = doc.add_paragraph()
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            heading.paragraph_format.space_after = Pt(8)
            heading_run = heading.add_run("PHOTO INDEX")
            heading_run.bold = True
            heading_run.font.size = Pt(20)
            
            # Calculate rows
            count = min(images_per_page, len(image_files) - start)
            rows_needed = (count + images_per_row - 1) // images_per_row
            
            table = doc.add_table(rows=rows_needed * 2, cols=images_per_row)
            table.autofit = True
            
            for i in range(count):
                global_index = start + i
                row = (i // images_per_row) * 2
                col = i % images_per_row
                img_path = os.path.join(source_folder, image_files[global_index])
                
                # Image cell
                img_cell = table.cell(row, col)
                img_para = img_cell.paragraphs[0]
                img_para.paragraph_format.space_after = Pt(0)
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = img_para.add_run()
                run.add_picture(img_path, width=image_width, height=image_height)
                
                # Caption cell
                cap_cell = table.cell(row + 1, col)
                cap_para = cap_cell.paragraphs[0]
                cap_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cap_para.paragraph_format.space_before = Pt(7)
                cap_run = cap_para.add_run(f"Photo {photo_counter:03d}")
                cap_run.font.size = Pt(12)
                photo_counter += 1
            
            if start + images_per_page < len(image_files):
                doc.add_page_break()
        
        doc.save(out_path)
        return out_path
    
    def generate_complete_report(
        self,
        record: PropertyRecord,
        middle_doc_path: str,
        photos_folder: str,
        output_folder: str
    ) -> str:
        """Full workflow - combines all documents"""
        timestamp = int(time.time())
        temp_dir = tempfile.gettempdir()
        
        # 1. Force landscape on middle doc
        middle_landscape = os.path.join(temp_dir, f"middle_land_{timestamp}.docx")
        self.force_docx_to_landscape_and_save(middle_doc_path, middle_landscape)
        
        # 2. Generate template
        template_path = os.path.join(temp_dir, f"template_{timestamp}.docx")
        self.generate_template_docx(record, template_path)
        
        # 3. Generate photo index
        photos_path = os.path.join(temp_dir, f"photos_{timestamp}.docx")
        self.build_photo_index_docx(photos_folder, photos_path)
        
        # 4. Merge documents
        final_path = os.path.join(
            output_folder,
            f"final_{record.id}_{record.client.replace(' ', '_')}_{timestamp}.docx"
        )
        
        # Use docxcompose to merge
        master = Document(template_path)
        self.set_doc_landscape(master)
        composer = Composer(master)
        
        # Append middle doc
        middle_doc = Document(middle_landscape)
        self.set_doc_landscape(middle_doc)
        composer.append(middle_doc)
        
        # Append photos
        photos_doc = Document(photos_path)
        self.set_doc_landscape(photos_doc)
        composer.append(photos_doc)
        
        composer.save(final_path)
        
        return final_path
