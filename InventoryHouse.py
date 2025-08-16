import sqlite3
from datetime import datetime
import tkinter as tk
from tkinter import messagebox, ttk, filedialog
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Cm, Pt
from docx.shared import Inches, Pt, Mm, RGBColor, Cm
from docx.enum.section import WD_ORIENT as WDO
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WDPA
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from lxml import etree
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import zipfile, shutil
import os
import subprocess, platform
import sys

# === CONFIGURATION ===
CLERK_OPTIONS = ["Tom Tyrrel", "Kevin Crack", "Bill West"]
STATUS_OPTIONS = ["Inspected", "Audio Recorded"]

# Determine base path (works for exe and script)
if getattr(sys, 'frozen', False):
    BASE_PATH = os.path.dirname(sys.executable)
else:
    BASE_PATH = os.path.dirname(os.path.abspath(__file__))

DB_PATH = os.path.join(BASE_PATH, "inventory.db")

# === DATABASE FUNCTIONS (SQLite) ===
def connect_db():
    conn = sqlite3.connect(DB_PATH)
    return conn

def create_table():
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS property_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        date TEXT,
        clerk TEXT,
        property_address TEXT,
        client TEXT,
        inv_type TEXT,
        status TEXT
    )
    """)
    conn.commit()
    conn.close()

# === APP FUNCTIONS ===
def add_record_popup():
    popup = tk.Toplevel(root)
    popup.title("Add New Record")

    tk.Label(popup, text="Clerk").grid(row=0, column=0)
    clerk_cb = ttk.Combobox(popup, values=CLERK_OPTIONS, state="readonly")
    clerk_cb.grid(row=0, column=1)

    tk.Label(popup, text="Address").grid(row=1, column=0)
    addr_entry = tk.Entry(popup)
    addr_entry.grid(row=1, column=1)

    tk.Label(popup, text="Client").grid(row=2, column=0)
    client_entry = tk.Entry(popup)
    client_entry.grid(row=2, column=1)

    tk.Label(popup, text="Inventory Type").grid(row=3, column=0)
    inv_entry = tk.Entry(popup)
    inv_entry.grid(row=3, column=1)

    tk.Label(popup, text="Status").grid(row=4, column=0)
    status_cb = ttk.Combobox(popup, values=STATUS_OPTIONS, state="readonly")
    status_cb.grid(row=4, column=1)

    def submit():
        values = [clerk_cb.get(), addr_entry.get(), client_entry.get(), inv_entry.get(), status_cb.get()]
        if not all(values):
            messagebox.showwarning("Input Error", "Please fill in all fields.")
            return
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
        INSERT INTO property_records (date, clerk, property_address, client, inv_type, status)
        VALUES (?, ?, ?, ?, ?, ?)
        """, (datetime.now().strftime("%Y-%m-%d"), *values))
        conn.commit()
        conn.close()
        popup.destroy()
        fetch_all()
        messagebox.showinfo("Success", "Record added successfully!")

    tk.Button(popup, text="Add Record", command=submit).grid(row=5, columnspan=2, pady=10)

def search_popup():
    popup = tk.Toplevel(root)
    popup.title("Search Records")

    tk.Label(popup, text="Client").grid(row=0, column=0)
    client_entry = tk.Entry(popup)
    client_entry.grid(row=0, column=1)

    tk.Label(popup, text="Clerk").grid(row=1, column=0)
    clerk_cb = ttk.Combobox(popup, values=CLERK_OPTIONS, state="readonly")
    clerk_cb.grid(row=1, column=1)

    tk.Label(popup, text="Address").grid(row=2, column=0)
    address_entry = tk.Entry(popup)
    address_entry.grid(row=2, column=1)

    def search():
        client = client_entry.get()
        clerk = clerk_cb.get()
        address = address_entry.get()
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
        SELECT * FROM property_records
        WHERE client LIKE ? AND clerk LIKE ? AND property_address LIKE ?
        """, (f"%{client}%", f"%{clerk}%", f"%{address}%"))
        rows = cursor.fetchall()
        conn.close()
        formatted = [(i+1, row[0], datetime.strptime(row[1], "%Y-%m-%d").strftime("%d-%m-%Y"), *row[2:]) for i, row in enumerate(rows)]
        update_tree(formatted)
        btn_clear_filters.config(state="normal")
        popup.destroy()

    tk.Button(popup, text="Search", command=search).grid(row=3, columnspan=2, pady=10)

def fetch_all():
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM property_records")
    rows = cursor.fetchall()
    conn.close()
    formatted = [(i+1, row[0], datetime.strptime(row[1], "%Y-%m-%d").strftime("%d-%m-%Y"), *row[2:]) for i, row in enumerate(rows)]
    update_tree(formatted)

def delete_record(record_id):
    if messagebox.askyesno("Confirm", "Delete this record?"):
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM property_records WHERE id = ?", (record_id,))
        conn.commit()
        conn.close()
        fetch_all()
        messagebox.showinfo("Deleted", "Record deleted successfully.")

def open_edit_popup(record):
    popup = tk.Toplevel(root)
    popup.title("Edit Record")

    is_completed = record[7].strip().lower() == "completed"

    tk.Label(popup, text="Clerk").grid(row=0, column=0, padx=5, pady=5)
    clerk_cb = ttk.Combobox(popup, values=CLERK_OPTIONS, state="readonly")
    clerk_cb.set(record[3])
    clerk_cb.grid(row=0, column=1, padx=5, pady=5)

    tk.Label(popup, text="Address").grid(row=1, column=0, padx=5, pady=5)
    address_entry = tk.Entry(popup)
    address_entry.insert(0, record[4])
    address_entry.grid(row=1, column=1, padx=5, pady=5)

    tk.Label(popup, text="Client").grid(row=2, column=0, padx=5, pady=5)
    client_entry = tk.Entry(popup)
    client_entry.insert(0, record[5])
    client_entry.grid(row=2, column=1, padx=5, pady=5)

    tk.Label(popup, text="Inventory Type").grid(row=3, column=0, padx=5, pady=5)
    inv_entry = tk.Entry(popup)
    inv_entry.insert(0, record[6])
    inv_entry.grid(row=3, column=1, padx=5, pady=5)

    tk.Label(popup, text="Status").grid(row=4, column=0, padx=5, pady=5)
    status_cb = ttk.Combobox(popup, values=STATUS_OPTIONS, state="readonly" if not is_completed else "disabled")
    status_cb.set(record[7])
    status_cb.grid(row=4, column=1, padx=5, pady=5)

    def save_changes():
        clerk = clerk_cb.get()
        address = address_entry.get().strip()
        client = client_entry.get().strip()
        inv_type = inv_entry.get().strip()
        status = record[7] if is_completed else status_cb.get()

        if not all([clerk, address, client, inv_type, status]):
            messagebox.showwarning("Input Error", "Please fill in all fields.")
            return

        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE property_records 
            SET clerk = ?, property_address = ?, client = ?, inv_type = ?, status = ?
            WHERE id = ?
        """, (clerk, address, client, inv_type, status, record[1]))
        conn.commit()
        conn.close()

        popup.destroy()
        fetch_all()
        messagebox.showinfo("Updated", "Record updated successfully.")

    tk.Button(popup, text="Save", command=save_changes).grid(row=5, columnspan=2, pady=10)

def generate_template(address, client, date):
    LOGO_PATH = os.path.join(BASE_PATH, "image 1.png")
    output_file = os.path.join(BASE_PATH, "template.docx")

    # Step 1: Build DOCX
    doc = Document()
    for section in doc.sections:
        section.orientation = WDO.LANDSCAPE
        section.page_width, section.page_height = Mm(297), Mm(210)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p = doc.add_paragraph()
    p.alignment = WDPA.LEFT
    p.add_run().add_picture(LOGO_PATH, width=Cm(12.38), height=Cm(2.9))

    # Substituted placeholders
    for text in [
        f"Property Address: {address}",
        f"On behalf of:     {client}",
        f"Date:             {date}"
    ]:
        para = doc.add_paragraph(text)
        para.runs[0].bold = True

    ph = doc.add_paragraph("Additional Notes")
    ph.runs[0].bold = True

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(25.46)
    tbl.rows[0].height = Cm(5.19)

    tbl_el = tbl._tbl
    tbl_pr = tbl_el.tblPr or OxmlElement('w:tblPr')
    tbl_el.insert(0, tbl_pr)
    tbl_borders = parse_xml(
        r'<w:tblBorders %s>'
        r'<w:top w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:left w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:right w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="auto"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    for old in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(old)
    tbl_pr.append(tbl_borders)

    cell = tbl.rows[0].cells[0]
    cell.paragraphs[0].text = "[NOTES_SDT_PARAGRAPH]"

    # Footer info
    doc.add_paragraph()
    f1 = doc.add_paragraph()
    run1 = f1.add_run("Inventory House ")
    run1.font.color.rgb = RGBColor(255, 0, 0)
    run1.font.size = Pt(12)
    run1.bold = True
    run2 = f1.add_run("T: 08700 336969 ")
    run2.font.size = Pt(12)

    email = "info@inventoryhouse.co.uk"
    website = "www.inventoryhouse.co.uk"

    def add_hyperlink(paragraph, url, text, color, bold, size):
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
        )
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
            rPr.append(c)
        if bold:
            rPr.append(OxmlElement('w:b'))
        if size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(size.pt * 2)))
            rPr.append(sz)
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    add_hyperlink(f1, f"mailto:{email}", email, RGBColor(0, 0, 255), False, Pt(12))
    f1.add_run("  ")
    add_hyperlink(f1, f"http://{website}", website, RGBColor(0, 0, 255), False, Pt(12))

    f2 = doc.add_paragraph(
        "Head Office: 3 County Gate London SE9 3UB.\n"
        "Inventory House Limited. Registered in England & Wales Company No. 5250554"
    )
    f2.runs[0].font.size = Pt(12)
    f2.runs[0].bold = True

    doc.save(output_file)
    return output_file

# paste_photos() stays the same except DB update changes:
def paste_photos(record_id):
    audio_transcription = filedialog.askopenfilename(title="Select audio transcribed Word file", filetypes=[
        ("Word Documents", "*.docx;*.doc"),
        ("All Files", "*.*")])
    if not audio_transcription:
        messagebox.showinfo("Invalid", "Choose a file of format .doc, .docx")
        return
    source_folder = filedialog.askdirectory(title="Select Image Folder")
    if not source_folder:
        messagebox.showinfo("Invalid", "No such directory")
        return
    try:
        output_file = os.path.join(BASE_PATH, "Photo_gallery.docx")
        # Get record info
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("SELECT property_address, client, date FROM property_records WHERE id = ?", (record_id,))
        addr, client, date_val = cursor.fetchone()
        conn.close()

        template_path = generate_template(addr, client, date_val)
        master = Document(template_path)

        images_per_page = 8
        images_per_row = 4
        image_width = Cm(5.85)
        image_height = Cm(6)

        master.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        composer = Composer(master)

        if os.path.exists(audio_transcription):
            middle_doc = Document(audio_transcription)
            composer.append(middle_doc)

        temp_merged_path = os.path.join(BASE_PATH, "temp_merged.docx")
        composer.save(temp_merged_path)

        doc = Document(temp_merged_path)
        #doc.add_page_break()

        image_files = sorted([
            f for f in os.listdir(source_folder)
            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
        ])

        photo_counter = 1
        for start in range(0, len(image_files), images_per_page):
            heading = doc.add_paragraph()
            heading.alignment = WDPA.LEFT
            heading.paragraph_format.space_after = Pt(8)
            heading_run = heading.add_run("PHOTO INDEX")
            heading_run.bold = True
            heading_run.font.size = Pt(20)

            table = doc.add_table(rows=4, cols=images_per_row)
            table.autofit = True
            for i in range(images_per_page):
                idx = start + i
                if idx >= len(image_files):
                    break
                row = (i // images_per_row) * 2
                col = i % images_per_row
                img_path = os.path.join(source_folder, image_files[idx])
                img_cell = table.cell(row, col)
                img_para = img_cell.paragraphs[0]
                img_para.paragraph_format.space_after = Pt(0)
                img_para.alignment = WDPA.CENTER
                run = img_para.add_run()
                run.add_picture(img_path, width=image_width, height=image_height)

                caption_cell = table.cell(row + 1, col)
                caption_para = caption_cell.paragraphs[0]
                caption_para.alignment = WDPA.CENTER
                caption_para.paragraph_format.space_before = Pt(7)
                caption_para.paragraph_format.space_after = Pt(7)
                caption_run = caption_para.add_run(f"Photo {photo_counter:03d}")
                caption_run.font.size = Pt(12)
                photo_counter += 1

        if os.path.exists(temp_merged_path):
            os.remove(temp_merged_path)
        if os.path.exists(output_file):
            os.remove(output_file)
        doc.save(output_file)

        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("UPDATE property_records SET status = ? WHERE id = ?", ("Completed", record_id))
        conn.commit()
        conn.close()
        fetch_all()
        messagebox.showinfo("Success", f"Photo document saved at:\n{output_file}\nStatus updated.")
        if platform.system() == "Windows":
            os.startfile(output_file)
        elif platform.system() == "Darwin":
            subprocess.call(["open", output_file])
        else:
            subprocess.call(["xdg-open", output_file])
    except Exception as e:
        messagebox.showerror("Error", str(e))

def clear_filters():
    fetch_all()
    btn_clear_filters.config(state="disabled")

# === GUI SETUP ===
root = tk.Tk()
root.title("InventoryHouse")
root.configure(bg="#9ECAD6")
root.geometry("1150x700")

img = Image.open(os.path.join(BASE_PATH, "logo.png"))
photo = ImageTk.PhotoImage(img)
title_label = tk.Label(root, bg="#9ECAD6", image=photo)
title_label.image = photo
title_label.pack(pady=10)

style = ttk.Style()
style.theme_use('clam')
style.configure("Treeview.Heading", background="#98A1BC", foreground="#555879", font=('Segoe UI', 15, 'bold'))

tree_frame = tk.Frame(root)
tree_frame.pack(fill="both", expand=True, padx=10, pady=5)
tree_y_scroll = tk.Scrollbar(tree_frame, orient="vertical")
tree_y_scroll.pack(side="right", fill="y")
tree_x_scroll = tk.Scrollbar(tree_frame, orient="horizontal", command=lambda *args: tree.xview(*args))
tree_x_scroll.pack(side="bottom", fill="x")

tree = ttk.Treeview(tree_frame, columns=("No", "ID", "Date", "Clerk", "Address", "Client", "Type", "Status"),
                    show="headings", yscrollcommand=tree_y_scroll.set, xscrollcommand=tree_x_scroll.set)
tree.tag_configure('oddrow', background='white', font=('Segoe UI', 12))
tree.tag_configure('evenrow', background='lightgray', font=('Segoe UI', 12))

def update_tree(rows):
    for item in tree.get_children():
        tree.delete(item)
    for i, row in enumerate(rows):
        tag = 'evenrow' if i % 2 == 0 else 'oddrow'
        tree.insert("", tk.END, values=row, tags=(tag,))
    reset_action_buttons()

tree_y_scroll.config(command=tree.yview)
tree_x_scroll.config(command=tree.xview)
for col in tree["columns"]:
    tree.heading(col, text=col)
    if col == "ID":
        tree.column(col, width=0, stretch=False)
    else:
        tree.column(col, width=120, anchor="center")
tree.pack(fill="both", expand=True)
tree.column("No", width=30, anchor="center")
tree.column("ID", width=0, stretch=False)
tree.column("Date", width=70, anchor="center")
tree.column("Clerk", width=80, anchor="center")
tree.column("Address", width=350, anchor="center", stretch=True)
tree.column("Client", width=100, anchor="center")
tree.column("Type", width=70, anchor="center")
tree.column("Status", width=90, anchor="center")
tree.bind("<<TreeviewSelect>>", lambda e: on_row_select(e))

action_frame = tk.LabelFrame(root, text="Table Actions", bg="#DFF6FF", fg="#3C5B6F", font=("Arial", 17, "bold"))
action_frame.pack(fill="x", padx=10, pady=5)
tk.Button(action_frame, text="Add Record", bg="#748DAE", fg="#2C3E50", font=("Arial", 14, "bold"), command=add_record_popup).pack(side="left", padx=5)
tk.Button(action_frame, text="Search", bg="#748DAE", fg="#2C3E50", font=("Arial", 14, "bold"), command=search_popup).pack(side="left", padx=5)
btn_clear_filters = tk.Button(action_frame, text="Clear Filters", state="disabled", bg="#748DAE", fg="#2C3E50", font=("Arial", 14, "bold"), command=lambda: clear_filters())
btn_clear_filters.pack(side="left", padx=5)

row_action_frame = tk.LabelFrame(root, text="Selected Row Actions", bg="#DFF6FF", fg="#3C5B6F", font=("Arial", 17, "bold"))
row_action_frame.pack(fill="x", padx=10, pady=5)
btn_edit = tk.Button(row_action_frame, text="Edit", state="disabled", bg="#748DAE", fg="#2C3E50", font=("Arial", 14, "bold"))
btn_edit.pack(side="left", padx=5)
btn_delete = tk.Button(row_action_frame, text="Delete", state="disabled", bg="#748DAE", fg="#2C3E50", font=("Arial", 14, "bold"))
btn_delete.pack(side="left", padx=5)
btn_photos = tk.Button(row_action_frame, text="Paste Photos", state="disabled", bg="#748DAE", fg="#2C3E50", font=("Arial", 14, "bold"))
btn_photos.pack(side="left", padx=5)

selected_record = None
def on_row_select(event):
    global selected_record
    selected_item = tree.selection()
    if selected_item:
        selected_record = tree.item(selected_item[0])['values']
        btn_edit.config(state="normal", command=lambda: open_edit_popup(selected_record))
        if selected_record[7].lower() == "completed":
            btn_photos.config(state="disabled")
        else:
            btn_photos.config(state="normal", command=lambda: paste_photos(selected_record[1]))
        btn_delete.config(state="normal", command=lambda: delete_record(selected_record[1]))
    else:
        reset_action_buttons()

def reset_action_buttons():
    btn_edit.config(state="disabled")
    btn_delete.config(state="disabled")
    btn_photos.config(state="disabled")

# Start App
create_table()
fetch_all()
root.mainloop()
